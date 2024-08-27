from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application,CommandHandler, MessageHandler,filters,ContextTypes,CallbackQueryHandler
from typing import Final
import requests
from docx import Document
import csv
TOKEN:Final="7201982465:AAEN9xfa3bsI_bLvLI9S7WtYYescpJxrjfk"
BOT_USERNAME:Final="@Orimi_BookieBot"
m=0
user=""
def searchbook(typ,query,ma):
    url="https://www.googleapis.com/books/v1/volumes"
    params={
        'q': f"{typ}:{query}",
        "key":"AIzaSyDSmOyU8osKNYRRCbdHOY6Vz69jBZ5-sfg",
        'maxResults': ma
        }
    re=requests.get(url,params=params)
    if re.status_code==200:
        return(re.json())
    else:
        re.raise_for_status()

async def start_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Hi There I'm Page Pal .Im Here To Help You Choose Your Next Fav Book!!!!type /help To View My Commands")
async def help_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Hi There I'm Page Pal .Im Here To Help You Choose Your Next Fav Book!!!!\nIf Your New Here Then Here Are My Commands\n /start To Start \n/book To Get A List Of Books Based On Genre Of Your Choice\n /preview To Get Sneak Peak On What Could Be Your Next Read \n/list To See And Manage Your Reading List\n/help To Review My Commands Again")
def makecsv(books_json):
    l=[["Title","Author","Genre","Published Date","Descriprion","Preveiw Link"]]
    with open("books.csv","w",newline="", encoding="utf-8") as f:
        csvw=csv.writer(f,delimiter=",")
        for item in books_json.get('items', []):
            volume_info = item.get('volumeInfo', {})
            title = volume_info.get('title', 'No title available')
            authors = volume_info.get('authors', 'No authors available')
            published_date = volume_info.get('publishedDate', 'No date available')
            description = volume_info.get('description', 'No description available')
            genre=volume_info.get("categories","no genre")
            preview=volume_info.get("previewLink","none")
            r=[title,authors,genre,published_date,description,preview]
            l.append(r)    
        for g in l:
            csvw.writerow(g)
            
def handle_response(text:str)-> str:
    text.lower()
    global m
    if m==1:
        result=searchbook("subject",text,40)
        makecsv(result)
        return("Searching Books Under "+text+" Genre")
    if m==2:
        result=searchbook("title",text,1)
        for item in result.get('items', []):
            volume_info = item.get('volumeInfo', {})
            preview=volume_info.get("previewLink","none")
            title=volume_info.get("title","No Title Available")
        title.lower()
        return(preview)
    
def read(file_path):
    full_text = []
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return(full_text)

async def handle_message(update:Update,context:ContextTypes.DEFAULT_TYPE):
    global m,user
    message_type:str=update.message.chat.type
    text:str=update.message.text
    text=text.lower()
    if message_type=="group":
        if BOT_USERNAME in text:
            new_text:str=text.replace(BOT_USERNAME,"").strip()
            response:str=handle_response(new_text)
        else:
            return
    else:
        response:str=handle_response(text)
    if m==1:
        await update.message.reply_text(response)
        chat_id = update.effective_chat.id
        await update.message.reply_text("Here Is A List Of Books I Found Under The "+text+" Genre\nHope This Helps You ;)")
        await context.bot.send_document(chat_id,"books.csv")
        m=0
    if m==2:
        await update.message.reply_text("Sure Here Is The Link: "+response)
        await update.message.reply_text("If This is not one you were looking For Please Try Again Or Be More Specific Thx ;)")
        m=0
    if m==3:
        chat_id = update.effective_chat.id
        file_path=user.first_name+" "+user.last_name+"s reading List"+".docx"
        try:
            doc = Document(file_path)
        except:
            doc=Document()
            doc.save(file_path)
        p=doc.add_paragraph()
        p.add_run(text)
        doc.save(file_path)
        await update.message.reply_text("Succelfully Added To Reading List")
        m=0
    if m==4:
        file_path=user.first_name+" "+user.last_name+"s reading List"+".docx"
        try:
            doc = Document(file_path)
        except:
            doc=Document()
            doc.save(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        if text in full_text:
            nd = Document()
            for g in full_text:
                if g!= text:
                    nd.add_paragraph(g)
            nd.save(file_path)
            await update.message.reply_text("Book Has Been Removed From Your Reading List")
        else:
            await update.message.reply_text("Book Not Found")
        m=0

        
async def error(update:Update,context:ContextTypes.DEFAULT_TYPE):
    print(f"Update{update} caused error{context.error}")
async def preview_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    global m
    m=2
    await update.message.reply_text("Please Enter The Name Of The Book You Want The Preview Link Of...")
async def book_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    global m    
    m=1
    await update.message.reply_text("Please Tell Me What Genre Of You Are Looking For")
async def list_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    global m,user
    user = update.message.from_user
    chat_id = update.effective_chat.id
    menu=[[InlineKeyboardButton("Get Reading List",callback_data="a")],[InlineKeyboardButton("Add Book To Reading List",callback_data="b")],[InlineKeyboardButton("Delete Book From Reading List",callback_data="c")]]
    await context.bot.send_message(chat_id,reply_markup=InlineKeyboardMarkup(menu),text="Please Choose From The Following Tasks")
async def menu_handler(update:Update,context:ContextTypes.DEFAULT_TYPE):
    global m,user
    chat_id = update.effective_chat.id
    option=update.callback_query.data
    if option=="a":
        file_path=user.first_name+" "+user.last_name+"s reading List"+".docx"
        try:
            doc = Document(file_path)
        except:
            doc=Document()
            doc.save(file_path)
        await context.bot.send_document(chat_id,file_path)
        m=0
    elif option=="b":
        m=3
        await context.bot.send_message(chat_id,"enter Book Name")
    elif option=="c":
        m=4
        await context.bot.send_message(chat_id,"enter Book Name")
if __name__== "__main__":
    app=Application.builder().token(TOKEN).build()
    app.add_handler(CommandHandler("start",start_command))
    app.add_handler(CommandHandler("help",help_command))
    app.add_handler(CommandHandler("book",book_command))
    app.add_handler(CommandHandler("preview",preview_command))
    app.add_handler(CommandHandler("list",list_command))
    app.add_handler(CallbackQueryHandler(menu_handler))
    app.add_handler(MessageHandler(filters.TEXT,handle_message))
    app.add_error_handler(error)
    app.run_polling(poll_interval=2)
    print("starting")

    
